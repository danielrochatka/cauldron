"""Unit tests for document extractors."""
from __future__ import annotations

import io

import pytest

from cauldron_ai_attachments.extractors import (
    DocxExtractor,
    ExtractionError,
    ExtractionResult,
    PdfExtractor,
    PlainTextExtractor,
    extract_document,
    get_extractor_for,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _make_docx(paragraphs: list[str], headings: list[str] | None = None) -> bytes:
    from docx import Document
    doc = Document()
    if headings:
        for h in headings:
            doc.add_heading(h, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_minimal_pdf(text: str) -> bytes:
    """Generate a minimal valid PDF bytes object with embedded text."""
    # Encode special chars for PDF string
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 50 750 Td ({safe}) Tj ET".encode("latin-1")
    stream_len = len(stream)

    # Build with hard-coded offsets that work for this template
    obj1 = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    obj2 = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    obj3 = (
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    obj4_header = (
        b"4 0 obj\n<< /Length "
        + str(stream_len).encode()
        + b" >>\nstream\n"
    )
    obj4 = obj4_header + stream + b"\nendstream\nendobj\n"
    obj5 = b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"

    header = b"%PDF-1.4\n"
    body = header + obj1 + obj2 + obj3 + obj4 + obj5

    # Compute real offsets
    o1 = len(header)
    o2 = o1 + len(obj1)
    o3 = o2 + len(obj2)
    o4 = o3 + len(obj3)
    o5 = o4 + len(obj4)
    xref_offset = len(body)

    xref = (
        b"xref\n0 6\n"
        + b"0000000000 65535 f \n"
        + f"{o1:010d} 00000 n \n".encode()
        + f"{o2:010d} 00000 n \n".encode()
        + f"{o3:010d} 00000 n \n".encode()
        + f"{o4:010d} 00000 n \n".encode()
        + f"{o5:010d} 00000 n \n".encode()
    )
    trailer = (
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF\n"
    )
    return body + xref + trailer


# ---------------------------------------------------------------------------
# PlainTextExtractor
# ---------------------------------------------------------------------------

class TestPlainTextExtractor:
    def setup_method(self):
        self.extractor = PlainTextExtractor()

    def test_can_extract_txt(self):
        assert self.extractor.can_extract("text/plain", "resume.txt")

    def test_can_extract_md(self):
        assert self.extractor.can_extract("text/plain", "readme.md")

    def test_cannot_extract_pdf(self):
        assert not self.extractor.can_extract("application/pdf", "file.pdf")

    def test_extract_basic_text(self):
        data = b"Jane Smith\nSoftware Engineer\nPython"
        result = self.extractor.extract(data)
        assert isinstance(result, ExtractionResult)
        assert "Jane Smith" in result.text
        assert result.word_count > 0
        assert result.extractor == "plain_text"

    def test_extract_markdown_headings(self):
        data = b"# Skills\n\nPython, Django\n\n## Experience\n\nFive years"
        result = self.extractor.extract(data)
        assert "Skills" in result.section_headings
        assert "Experience" in result.section_headings

    def test_truncates_long_text(self):
        data = ("word " * 20000).encode()
        result = self.extractor.extract(data, max_chars=100)
        assert result.truncated
        assert len(result.text) <= 100

    def test_handles_utf8(self):
        data = "Héllo wörld — résumé".encode("utf-8")
        result = self.extractor.extract(data)
        assert "résumé" in result.text


# ---------------------------------------------------------------------------
# PdfExtractor
# ---------------------------------------------------------------------------

class TestPdfExtractor:
    def setup_method(self):
        self.extractor = PdfExtractor()

    def test_can_extract_pdf(self):
        assert self.extractor.can_extract("application/pdf", "cv.pdf")
        assert self.extractor.can_extract("text/plain", "file.pdf")  # extension wins

    def test_cannot_extract_docx(self):
        assert not self.extractor.can_extract("text/plain", "file.docx")

    def test_extract_minimal_pdf(self):
        pdf_bytes = _make_minimal_pdf("Alice Engineer Python Django")
        result = self.extractor.extract(pdf_bytes)
        assert isinstance(result, ExtractionResult)
        assert result.extractor == "pdf"
        assert result.page_count >= 1

    def test_corrupt_pdf_raises_extraction_error(self):
        bad = b"%PDF-1.4\nThis is not a real PDF."
        with pytest.raises(ExtractionError, match="PDF could not be read"):
            self.extractor.extract(bad)

    def test_non_pdf_bytes_raises_extraction_error(self):
        with pytest.raises(ExtractionError):
            self.extractor.extract(b"This is just plain text")


# ---------------------------------------------------------------------------
# DocxExtractor
# ---------------------------------------------------------------------------

class TestDocxExtractor:
    def setup_method(self):
        self.extractor = DocxExtractor()

    def test_can_extract_docx(self):
        assert self.extractor.can_extract(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "resume.docx",
        )
        assert self.extractor.can_extract("text/plain", "file.docx")

    def test_cannot_extract_txt(self):
        assert not self.extractor.can_extract("text/plain", "file.txt")

    def test_extract_paragraphs(self):
        docx = _make_docx(["Bob Lee", "Senior Developer", "TypeScript, React"])
        result = self.extractor.extract(docx)
        assert "Bob Lee" in result.text
        assert result.word_count > 0
        assert result.extractor == "docx"

    def test_extract_headings(self):
        from docx import Document
        doc = Document()
        doc.add_heading("Skills", level=1)
        doc.add_paragraph("Python, Go")
        doc.add_heading("Experience", level=2)
        buf = io.BytesIO()
        doc.save(buf)
        result = self.extractor.extract(buf.getvalue())
        assert "Skills" in result.section_headings
        assert "Experience" in result.section_headings

    def test_truncates_long_docx(self):
        docx = _make_docx(["word " * 1000])
        result = self.extractor.extract(docx, max_chars=50)
        assert result.truncated
        assert len(result.text) <= 50

    def test_corrupt_docx_raises_extraction_error(self):
        with pytest.raises(ExtractionError, match="DOCX could not be read"):
            self.extractor.extract(b"not a docx file")


# ---------------------------------------------------------------------------
# extract_document dispatcher
# ---------------------------------------------------------------------------

def test_extract_document_dispatches_to_pdf():
    pdf = _make_minimal_pdf("Dispatch test")
    result = extract_document(pdf, "application/pdf", "test.pdf")
    assert result.extractor == "pdf"


def test_extract_document_dispatches_to_docx():
    docx = _make_docx(["Dispatch test"])
    result = extract_document(
        docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "test.docx",
    )
    assert result.extractor == "docx"


def test_extract_document_dispatches_to_plain():
    result = extract_document(b"Hello world", "text/plain", "note.txt")
    assert result.extractor == "plain_text"


def test_extract_document_raises_for_unsupported():
    with pytest.raises(ExtractionError, match="No extractor"):
        extract_document(b"\xff\xd8\xff", "image/jpeg", "photo.jpg")


def test_get_extractor_for_returns_none_on_unknown():
    assert get_extractor_for("image/png", "photo.png") is None
